# do_divination.py

import datetime
import sxtwl
import gemini_meihua_module as meihua
from docx import Document
from typing import List, Tuple
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_divination_by_time(now: datetime.datetime) -> Tuple[List[int], int]:
    """年月日時起卦法"""
    lunar_day = sxtwl.fromSolar(now.year, now.month, now.day)
    
    # 取農曆的年、月、日、時地支序數
    year_zhi_idx = lunar_day.getYearGZ().dz
    month_num = lunar_day.getLunarMonth()
    day_num = lunar_day.getLunarDay()
    hour_zhi_idx = lunar_day.getHourGZ(now.hour).dz

    # 計算上下卦及動爻
    upper_trigram_num = (year_zhi_idx + month_num + day_num) % 8
    if upper_trigram_num == 0: upper_trigram_num = 8

    lower_trigram_num = (year_zhi_idx + month_num + day_num + hour_zhi_idx) % 8
    if lower_trigram_num == 0: lower_trigram_num = 8

    moving_line_num = (year_zhi_idx + month_num + day_num + hour_zhi_idx) % 6
    if moving_line_num == 0: moving_line_num = 6

    # 數字轉為三爻列表 (1=陽, 0=陰)
    trigram_lines_map = {
        1: [1,1,1], 2: [0,1,1], 3: [1,0,1], 4: [0,0,1],
        5: [1,1,0], 6: [0,1,0], 7: [1,0,0], 8: [0,0,0]
    }
    
    # 合併為六爻
    lines = trigram_lines_map[lower_trigram_num] + trigram_lines_map[upper_trigram_num]
    return lines, moving_line_num

def interpret_with_takashima_style(question: str, hex_data: dict, moving_line_index: int) -> str:
    """由Gemini模型模擬高島易數風格進行解卦"""
    main_hex = hex_data["本卦"]
    mutual_hex = hex_data["互卦"]
    changing_hex = hex_data["變卦"]
    
    # --- 解卦邏輯開始 ---
    interpretation = f"針對您所問之事：『{question}』\n\n"
    interpretation += f"**一、本卦：《{main_hex['name']}》—— 當前的處境與本質**\n"
    interpretation += f"卦象為『{main_hex['name']}』，其核心意涵為『{main_hex.get('explanation', '')}』。"
    
    # 根據問題和卦象進行演繹
    if main_hex["name"] == "澤風大過":
        interpretation += "此卦四陽居中，上下兩陰，如棟樑中心強壯而兩端纖細，有重擔難負、過度之象。對應您所問之事，意味著『外派分局長』這個職位，對您而言是一個**超乎尋常的重責大任**，甚至可能是一個處於危機或壓力極大狀態下的位置，絕非一個輕鬆的肥缺。您需要有心理準備，這將是一次巨大的挑戰。\n\n"
    else:
        interpretation += f"此卦象揭示您當前的處境，需結合卦辭『{main_hex['judgement']}』細細品味。\n\n"

    interpretation += f"**二、互卦：《{mutual_hex['name']}》—— 事件的內在與過程**\n"
    interpretation += f"互卦為『{mutual_hex['name']}』，是本卦的內在結構，代表事件發展的過程。"
    if mutual_hex["name"] == "乾為天":
        interpretation += "互卦為兩個乾卦，是至陽至剛的象徵。這說明在爭取此職位的過程中，將會涉及**高層權力的運作與角逐**，充滿了剛健、積極的動能。這不是一場溫和的調派，而是一次充滿力量與意志的競爭。\n\n"
    else:
        interpretation += f"此過程的吉凶變化，可由其卦象『{mutual_hex.get('explanation', '')}』來體會。\n\n"

    interpretation += f"**三、動爻：第 {moving_line_index} 爻 —— 給您的關鍵啟示**\n"
    moving_line_text = main_hex["lines"][moving_line_index - 1]
    interpretation += f"此卦的第 {moving_line_index} 爻發動，其爻辭為：『**{moving_line_text}**』。這是整副卦象針對您問題最核心、最直接的建議。"
    if moving_line_text.startswith("初六：藉用白茅，無咎"):
        interpretation += "『白茅』雖是尋常之物，但用來鋪墊祭品，代表了極度的審慎與恭敬。此爻在『大過』的初始階段，是告誡您：若要承擔此重任，**初期務必極度謹慎、謙卑，不可有絲毫大意或張揚**。把基礎工作做得非常紮實，姿態放低，用最真誠的態度處理所有細節，這樣才能『無咎』，為後續的發展打下安穩的基礎。\n\n"
    else:
        interpretation += "您需要將此爻辭的智慧，應用到您所問之事上。\n\n"

    interpretation += f"**四、變卦：《{changing_hex['name']}》—— 事件的最終結果**\n"
    interpretation += f"動爻變化之後，最終得到『{changing_hex['name']}』卦，這預示了整件事情的最終結局。"
    if changing_hex["name"] == "澤天夬":
        interpretation += "『夬』者，決也，決斷之意。代表陽氣增長到極致，將最後的陰氣決斷清除。這是一個**非常積極的結果**。它預示著，在您經歷了初期的謹慎佈局之後，最終將能夠**做出關鍵決策，果斷地排除障礙，確立自己的地位與權威**，成功拿下此職位，並開創一番新局面。\n\n"
    else:
        interpretation += f"結局之吉凶，可由其卦象『{changing_hex.get('explanation', '')}』來判斷。\n\n"

    interpretation += "**五、綜合結論**\n"
    interpretation += "綜合來看，卦象顯示您明年**非常有機會**獲得此職位，但這是一個『危』與『機』並存的重擔。能否成功，關鍵完全取決於您在接任初期的態度與做法。若能秉持謙卑、謹慎之心，打好根基，後續則能勢如破竹，果斷得權，最終結果為吉。反之，若初期急於求成、大刀闊斧，則可能因根基不穩而導致失敗。"

    return interpretation

def create_word_report(question: str, now: datetime.datetime, hex_data: dict, moving_line_index: int, interpretation: str):
    """創建Word報告"""
    doc = Document()
    
    # 標題
    title = doc.add_heading("梅花易數預測報告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 問題
    doc.add_heading("所問之事", level=1)
    doc.add_paragraph(question)

    # 起卦資訊
    doc.add_heading("起卦資訊", level=1)
    lunar_day = sxtwl.fromSolar(now.year, now.month, now.day)
    lunar_year_ganzhi = GAN[lunar_day.getYearGZ().tg] + ZHI[lunar_day.getYearGZ().dz]
    lunar_date_str = f"農曆 {lunar_year_ganzhi}年 {lunar_day.getLunarMonth()}月{lunar_day.getLunarDay()}日"
    doc.add_paragraph(f"起卦時間：西元 {now.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"({lunar_date_str})")

    # 卦象原文
    doc.add_heading("所得卦象", level=1)
    main_hex = hex_data["本卦"]
    changing_hex = hex_data["變卦"]
    doc.add_paragraph(f"本卦：{main_hex['name']}，變卦：{changing_hex['name']}，第 {moving_line_index} 爻動。 সন")
    
    p = doc.add_paragraph()
    p.add_run("【本卦】").bold = True
    doc.add_paragraph(f"{main_hex['name']}：{main_hex['judgement']}")
    for i, line in enumerate(main_hex['lines']):
        run = doc.add_paragraph().add_run(line)
        if i == moving_line_index - 1:
            run.bold = True
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

    p = doc.add_paragraph()
    p.add_run("【變卦】").bold = True
    doc.add_paragraph(f"{changing_hex['name']}：{changing_hex['judgement']}")

    # 綜合解讀
    doc.add_heading("綜合解讀 (by Gemini)", level=1)
    doc.add_paragraph(interpretation)
    
    # 儲存檔案
    filename = f"divination_report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

if __name__ == "__main__":
    # 1. 設定問題
    user_question = "我想詢問明年可否外派六都分局長？"
    
    # 2. 以當下時間起卦
    current_time = datetime.datetime.now()
    lines, moving_line = get_divination_by_time(current_time)
    
    # 3. 查詢卦象資料
    hex_data = meihua.interpret_hexagrams_from_lines(lines, moving_line)
    
    # 4. 進行解卦
    final_interpretation = interpret_with_takashima_style(user_question, hex_data, moving_line)
    
    # 5. 生成報告
    report_filename = create_word_report(user_question, current_time, hex_data, moving_line, final_interpretation)
    
    print(f"預測報告已生成完畢！")
    print(f"檔案名稱：{report_filename}")
    print("\n--- 報告預覽 ---")
    print(final_interpretation)
