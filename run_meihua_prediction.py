# run_meihua_prediction.py (v7 - 數字起卦版)

import datetime
import random
import gemini_meihua_module as meihua
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Tuple
import os
import google.generativeai as genai

# --- API 設定 ---
try:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("錯誤：請先設定名為 GEMINI_API_KEY 的環境變數。")
    genai.configure(api_key=api_key)
except Exception as e:
    print(e)
    exit()

# --- 核心功能函式 ---

def get_divination_by_numbers() -> Tuple[List[int], int, Tuple[int, int, int]]:
    """數字起卦法"""
    # 隨機取得三個1-100之間的數
    num1 = random.randint(1, 100)
    num2 = random.randint(1, 100)
    num3 = random.randint(1, 100)

    # 換算上下卦 (先天八卦數: 乾一、兌二、離三、震四、巽五、坎六、艮七、坤八)
    # 第一數為上卦, 第二數為下卦
    upper_trigram_num = num1 % 8 or 8
    lower_trigram_num = num2 % 8 or 8
    
    # 換算變爻
    moving_line_num = num3 % 6 or 6

    trigram_lines_map = {
        1: [1,1,1], 2: [0,1,1], 3: [1,0,1], 4: [0,0,1], # 乾、兌、離、震
        5: [1,1,0], 6: [0,1,0], 7: [1,0,0], 8: [0,0,0]  # 巽、坎、艮、坤
    }
    
    # 六爻順序由下至上，故 list = lower + upper
    lines = trigram_lines_map[lower_trigram_num] + trigram_lines_map[upper_trigram_num]
    
    return lines, moving_line_num, (num1, num2, num3)

def generate_interpretation_prompt(question: str, numbers: Tuple[int, int, int], hex_data: dict, moving_line_index: int) -> str:
    """生成需要被解讀的完整資訊"""
    main_hex = hex_data.get("本卦", {})
    mutual_hex = hex_data.get("互卦", {})
    changing_hex = hex_data.get("變卦", {})
    
    prompt = f'''
請扮演一位精通《易經》與高島易數風格的解卦專家，為我分析以下卦象：

**1. 我的問題：**
{question}

**2. 起卦資訊：**
- 起卦方式：數字起卦
- 所用數字：{numbers[0]} (上卦), {numbers[1]} (下卦), {numbers[2]} (變爻)
- 動爻：第 {moving_line_index} 爻

**3. 卦象結果：**

*   **本卦：《{main_hex.get('name', '未知')}》**
    *   卦辭：{main_hex.get('judgement', '')}
    *   第 {moving_line_index} 爻爻辭：{main_hex.get('lines', [''] * 6)[moving_line_index - 1]}

*   **互卦：《{mutual_hex.get('name', '未知')}》**
    *   卦辭：{mutual_hex.get('judgement', '')}

*   **變卦：《{changing_hex.get('name', '未知')}》**
    *   卦辭：{changing_hex.get('judgement', '')}

請結合我的問題，對「本卦」、「互卦」、「動爻」、「變卦」的關聯與意義進行全面、深入的解讀，並提供具體的結論與建議。
'''
    return prompt

def call_gemini_api(prompt: str) -> str:
    """呼叫 Gemini API 進行解讀"""
    print("正在呼叫 Gemini API 進行解卦，請稍候...")
    try:
        model = genai.GenerativeModel('models/gemini-pro-latest')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"呼叫 Gemini API 時出錯：{e}"

def create_word_report(question: str, numbers: Tuple[int, int, int], hex_data: dict, moving_line_index: int, interpretation: str) -> str:
    """創建Word報告"""
    doc = Document()
    now = datetime.datetime.now()
    
    title = doc.add_heading("梅花易數預測報告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"報告時間：{now.strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_heading("所問之事", level=1)
    doc.add_paragraph(question)
    
    doc.add_heading("起卦資訊", level=1)
    doc.add_paragraph(f"起卦方式：數字起卦")
    doc.add_paragraph(f"所用數字：{numbers[0]} (上卦), {numbers[1]} (下卦), {numbers[2]} (變爻)")

    doc.add_heading("所得卦象", level=1)
    main_hex = hex_data.get("本卦", {})
    changing_hex = hex_data.get("變卦", {})
    doc.add_paragraph(f"本卦：{main_hex.get('name', '未知')}，變卦：{changing_hex.get('name', '未知')}，第 {moving_line_index} 爻動。")
    
    p = doc.add_paragraph()
    p.add_run("【本卦】").bold = True
    doc.add_paragraph(f"{main_hex.get('name', '未知')}：{main_hex.get('judgement', '')}")
    lines = main_hex.get('lines', [])
    if lines:
        for i, line in enumerate(lines):
            run = doc.add_paragraph().add_run(line)
            if i == moving_line_index - 1:
                run.bold = True
                
    p = doc.add_paragraph()
    p.add_run("【變卦】").bold = True
    doc.add_paragraph(f"{changing_hex.get('name', '未知')}：{changing_hex.get('judgement', '')}")
    
    doc.add_heading("綜合解讀 (by Gemini API)", level=1)
    doc.add_paragraph(interpretation)
    
    filename = f"divination_report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

if __name__ == "__main__":
    user_question = input("請輸入您想詢問的具體事由，然後按下 Enter：")
    
    if user_question:
        # 1. 改為數字起卦
        lines, moving_line, numbers = get_divination_by_numbers()
        print(f"本次隨機取數結果： 上卦={numbers[0]}, 下卦={numbers[1]}, 變爻={numbers[2]}")

        # 2. 取得卦象資料
        hex_data = meihua.interpret_hexagrams_from_lines(lines, moving_line)
        
        # 3. 生成給 AI 的 Prompt
        prompt_for_gemini = generate_interpretation_prompt(user_question, numbers, hex_data, moving_line)
        
        # 4. 呼叫 AI 解卦
        final_interpretation = call_gemini_api(prompt_for_gemini)
        
        print("解卦完成，正在生成 Word 報告...")
        
        # 5. 生成報告
        try:
            report_filename = create_word_report(user_question, numbers, hex_data, moving_line, final_interpretation)
            print(f"\n預測報告已生成完畢！")
            print(f"檔案名稱：{report_filename}")
        except Exception as e:
            print(f"生成Word報告時出錯：{e}")
        
        print("\n--- Gemini API 解讀預覽 ---")
        print(final_interpretation)